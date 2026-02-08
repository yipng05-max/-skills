#!/usr/bin/env python3
"""
Analyze a .docx sample paper and extract formatting/style rules as JSON.

Usage:
    python analyze_docx.py <input.docx> [output.json]

Dependencies:
    pip install python-docx
"""

import sys
import json
import re
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm, Inches, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_ORIENT
except ImportError:
    print("Error: python-docx is required. Install with: pip install python-docx")
    sys.exit(1)


def emu_to_cm(emu):
    """Convert EMU to centimeters."""
    if emu is None:
        return None
    return round(emu / 360000, 2)


def emu_to_pt(emu):
    """Convert EMU to points."""
    if emu is None:
        return None
    return round(emu / 12700, 1)


def spacing_to_value(spacing):
    """Convert line spacing to a readable value."""
    if spacing is None:
        return None
    # Line spacing in docx is stored in EMU or as a multiple (x240)
    return round(spacing / 12700, 1)


def get_alignment_str(alignment):
    """Convert alignment enum to string."""
    mapping = {
        WD_ALIGN_PARAGRAPH.LEFT: "left",
        WD_ALIGN_PARAGRAPH.CENTER: "center",
        WD_ALIGN_PARAGRAPH.RIGHT: "right",
        WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
        WD_ALIGN_PARAGRAPH.DISTRIBUTE: "distribute",
    }
    return mapping.get(alignment, "left")


def analyze_page_layout(doc):
    """Extract page layout settings from all sections."""
    sections = []
    for i, section in enumerate(doc.sections):
        sec_info = {
            "section_index": i,
            "page_width_cm": emu_to_cm(section.page_width),
            "page_height_cm": emu_to_cm(section.page_height),
            "orientation": "landscape" if section.orientation == WD_ORIENT.LANDSCAPE else "portrait",
            "margins": {
                "top_cm": emu_to_cm(section.top_margin),
                "bottom_cm": emu_to_cm(section.bottom_margin),
                "left_cm": emu_to_cm(section.left_margin),
                "right_cm": emu_to_cm(section.right_margin),
            },
            "gutter_cm": emu_to_cm(section.gutter),
            "header_distance_cm": emu_to_cm(section.header_distance),
            "footer_distance_cm": emu_to_cm(section.footer_distance),
        }
        # Detect approximate paper size
        w = sec_info["page_width_cm"] or 0
        h = sec_info["page_height_cm"] or 0
        if abs(w - 21.0) < 0.5 and abs(h - 29.7) < 0.5:
            sec_info["paper_size"] = "A4"
        elif abs(w - 21.59) < 0.5 and abs(h - 27.94) < 0.5:
            sec_info["paper_size"] = "Letter"
        else:
            sec_info["paper_size"] = "custom"

        # Header/footer content
        if section.header and section.header.paragraphs:
            sec_info["header_text"] = " | ".join(
                p.text.strip() for p in section.header.paragraphs if p.text.strip()
            )
        if section.footer and section.footer.paragraphs:
            sec_info["footer_text"] = " | ".join(
                p.text.strip() for p in section.footer.paragraphs if p.text.strip()
            )

        sections.append(sec_info)
    return sections


def analyze_paragraph_format(para):
    """Extract formatting details from a paragraph."""
    pf = para.paragraph_format
    info = {
        "text_preview": para.text[:80] if para.text else "",
        "style_name": para.style.name if para.style else None,
        "alignment": get_alignment_str(pf.alignment) if pf.alignment is not None else None,
        "line_spacing": None,
        "line_spacing_rule": None,
        "space_before_pt": emu_to_pt(pf.space_before) if pf.space_before else None,
        "space_after_pt": emu_to_pt(pf.space_after) if pf.space_after else None,
        "first_line_indent_cm": emu_to_cm(pf.first_line_indent) if pf.first_line_indent else None,
        "left_indent_cm": emu_to_cm(pf.left_indent) if pf.left_indent else None,
    }

    # Line spacing
    if pf.line_spacing is not None:
        if pf.line_spacing_rule is not None:
            rule_name = str(pf.line_spacing_rule).split(".")[-1] if pf.line_spacing_rule else None
            info["line_spacing_rule"] = rule_name
            if rule_name in ("EXACTLY", "AT_LEAST"):
                info["line_spacing"] = f"{emu_to_pt(pf.line_spacing)}pt"
            else:
                info["line_spacing"] = round(pf.line_spacing, 2)
        else:
            info["line_spacing"] = round(pf.line_spacing, 2)

    # Character formatting from runs
    fonts = []
    sizes = []
    bold_runs = 0
    italic_runs = 0
    total_runs = 0
    for run in para.runs:
        total_runs += 1
        if run.font.name:
            fonts.append(run.font.name)
        if run.font.size:
            sizes.append(emu_to_pt(run.font.size))
        if run.font.bold:
            bold_runs += 1
        if run.font.italic:
            italic_runs += 1

    if fonts:
        info["font_names"] = list(dict.fromkeys(fonts))  # unique, preserve order
    if sizes:
        info["font_sizes_pt"] = list(dict.fromkeys(sizes))
    if total_runs > 0:
        info["bold"] = bold_runs > total_runs * 0.5
        info["italic"] = italic_runs > total_runs * 0.5

    return info


def classify_paragraph(para):
    """Classify a paragraph by its role (title, heading, body, caption, reference, etc.)."""
    style_name = (para.style.name or "").lower()
    text = para.text.strip()

    if not text:
        return "empty"

    # Heading detection by style
    if "heading" in style_name or "title" in style_name:
        # Extract level
        match = re.search(r"(\d+)", style_name)
        if "title" in style_name and "heading" not in style_name:
            return "title"
        level = int(match.group(1)) if match else 1
        return f"heading_{level}"

    # Caption detection
    if style_name.startswith("caption") or re.match(
        r"^(figure|fig\.|table|表|图)\s*\d", text, re.IGNORECASE
    ):
        return "caption"

    # Reference/bibliography detection
    if style_name in ("bibliography", "references"):
        return "reference"
    if re.match(r"^\[?\d+\]?\s", text) and len(text) > 30:
        return "reference"

    # Abstract detection
    if "abstract" in style_name or text.lower().startswith("abstract") or text.startswith("摘要"):
        return "abstract"

    # Keywords
    if text.lower().startswith("keywords") or text.startswith("关键词"):
        return "keywords"

    return "body"


def analyze_styles(doc):
    """Analyze all named styles defined in the document."""
    styles_info = {}
    for style in doc.styles:
        if style.type is not None and hasattr(style, "font"):
            s = {
                "type": str(style.type).split(".")[-1],
                "base_style": style.base_style.name if style.base_style else None,
            }
            if style.font:
                if style.font.name:
                    s["font_name"] = style.font.name
                if style.font.size:
                    s["font_size_pt"] = emu_to_pt(style.font.size)
                if style.font.bold is not None:
                    s["bold"] = style.font.bold
                if style.font.italic is not None:
                    s["italic"] = style.font.italic
                if style.font.color and style.font.color.rgb:
                    s["color"] = str(style.font.color.rgb)
            if hasattr(style, "paragraph_format") and style.paragraph_format:
                pf = style.paragraph_format
                if pf.alignment is not None:
                    s["alignment"] = get_alignment_str(pf.alignment)
                if pf.line_spacing is not None:
                    s["line_spacing"] = round(pf.line_spacing, 2) if isinstance(pf.line_spacing, float) else emu_to_pt(pf.line_spacing)
                if pf.first_line_indent:
                    s["first_line_indent_cm"] = emu_to_cm(pf.first_line_indent)
                if pf.space_before:
                    s["space_before_pt"] = emu_to_pt(pf.space_before)
                if pf.space_after:
                    s["space_after_pt"] = emu_to_pt(pf.space_after)
            styles_info[style.name] = s
    return styles_info


def analyze_tables(doc):
    """Analyze table formatting."""
    tables_info = []
    for i, table in enumerate(doc.tables):
        t = {
            "table_index": i,
            "rows": len(table.rows),
            "columns": len(table.columns),
            "style": table.style.name if table.style else None,
        }
        # Check first row for header-like formatting
        if table.rows:
            first_row_bold = all(
                cell.paragraphs[0].runs[0].font.bold
                for cell in table.rows[0].cells
                if cell.paragraphs and cell.paragraphs[0].runs
            )
            t["first_row_bold"] = first_row_bold
        tables_info.append(t)
    return tables_info


def detect_citation_style(paragraphs):
    """Attempt to detect the citation style used in the document."""
    in_text_patterns = {
        "numeric_bracket": r"\[\d+\]",
        "numeric_superscript": r"\^\[\d+\]",  # rough heuristic
        "author_year_paren": r"\([A-Z][a-z]+(?:\s(?:and|&)\s[A-Z][a-z]+)?,?\s*\d{4}\)",
        "author_year_cn": r"[（(]\s*[\u4e00-\u9fff]+\s*[,，]\s*\d{4}\s*[)）]",
    }

    counts = {k: 0 for k in in_text_patterns}
    for para in paragraphs:
        text = para.text
        for name, pattern in in_text_patterns.items():
            counts[name] += len(re.findall(pattern, text))

    detected = max(counts, key=counts.get) if any(v > 0 for v in counts.values()) else "unknown"

    style_map = {
        "numeric_bracket": "numeric (e.g., GB/T 7714 numeric, IEEE, Vancouver)",
        "numeric_superscript": "numeric superscript (e.g., Vancouver)",
        "author_year_paren": "author-year (e.g., APA, Chicago author-date)",
        "author_year_cn": "author-year Chinese (e.g., GB/T 7714 author-year)",
    }

    return {
        "detected_style": style_map.get(detected, "unknown"),
        "citation_counts": counts,
    }


def analyze_document(docx_path):
    """Main analysis function."""
    doc = Document(docx_path)

    # Classify and group paragraphs
    classified = {}
    for para in doc.paragraphs:
        role = classify_paragraph(para)
        if role == "empty":
            continue
        if role not in classified:
            classified[role] = []
        classified[role].append(analyze_paragraph_format(para))

    # Build summary: pick representative formatting for each role
    formatting_summary = {}
    for role, paras in classified.items():
        if not paras:
            continue
        # Use the first non-empty paragraph as representative
        rep = paras[0]
        formatting_summary[role] = {
            "count": len(paras),
            "font_names": rep.get("font_names"),
            "font_sizes_pt": rep.get("font_sizes_pt"),
            "bold": rep.get("bold"),
            "italic": rep.get("italic"),
            "alignment": rep.get("alignment"),
            "line_spacing": rep.get("line_spacing"),
            "line_spacing_rule": rep.get("line_spacing_rule"),
            "first_line_indent_cm": rep.get("first_line_indent_cm"),
            "space_before_pt": rep.get("space_before_pt"),
            "space_after_pt": rep.get("space_after_pt"),
            "sample_text": rep.get("text_preview"),
        }

    result = {
        "source_file": str(docx_path),
        "page_layout": analyze_page_layout(doc),
        "defined_styles": analyze_styles(doc),
        "paragraph_formatting": formatting_summary,
        "table_formatting": analyze_tables(doc),
        "citation_style": detect_citation_style(doc.paragraphs),
    }

    return result


def main():
    if len(sys.argv) < 2:
        print(f"Usage: {sys.argv[0]} <input.docx> [output.json]")
        sys.exit(1)

    input_path = Path(sys.argv[1])
    if not input_path.exists():
        print(f"Error: File not found: {input_path}")
        sys.exit(1)

    output_path = Path(sys.argv[2]) if len(sys.argv) > 2 else input_path.with_suffix(".style-rules.json")

    print(f"Analyzing: {input_path}")
    result = analyze_document(input_path)

    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(result, f, ensure_ascii=False, indent=2)

    print(f"Style rules saved to: {output_path}")
    print(json.dumps(result, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
