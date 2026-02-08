#!/usr/bin/env python3
"""
Generate a Word (.docx) document for a literature review.

Usage:
    python generate_docx.py --input review.md --output literature_review.docx
    python generate_docx.py --input review.md --output literature_review.docx --title "论文标题"

Input: A markdown file with the literature review content. Expected structure:
    - Lines starting with ## are section headings (e.g., ## Literature Review, ## References)
    - Lines starting with ### are subsection headings
    - Other lines are body paragraphs
    - A blank line separates paragraphs

Output: A formatted .docx file with Times New Roman, proper heading styles, and
        hanging indent for references.
"""

import argparse
import re
import sys

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def set_normal_style(doc):
    """Configure the Normal style for body text."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    pf = style.paragraph_format
    pf.line_spacing = 2.0
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    pf.first_line_indent = Cm(0.75)
    # Set East Asian font
    rpr = style.element.find(qn("w:rPr"))
    if rpr is None:
        rpr = style.element.makeelement(qn("w:rPr"), {})
        style.element.append(rpr)
    ea = rpr.find(qn("w:rFonts"))
    if ea is None:
        ea = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(ea)
    ea.set(qn("w:eastAsia"), "宋体")


def add_heading_styled(doc, text, level):
    """Add a heading with Times New Roman font."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = None  # Remove default color
        rpr = run._element.find(qn("w:rPr"))
        if rpr is not None:
            ea = rpr.find(qn("w:rFonts"))
            if ea is None:
                ea = rpr.makeelement(qn("w:rFonts"), {})
                rpr.append(ea)
            ea.set(qn("w:eastAsia"), "黑体")
    heading.paragraph_format.space_before = Pt(12)
    heading.paragraph_format.space_after = Pt(6)
    heading.paragraph_format.first_line_indent = Cm(0)
    return heading


def add_reference_paragraph(doc, text):
    """Add a reference entry with hanging indent."""
    para = doc.add_paragraph(text)
    para.style = doc.styles["Normal"]
    pf = para.paragraph_format
    pf.first_line_indent = Cm(-1.27)  # Hanging indent
    pf.left_indent = Cm(1.27)
    pf.line_spacing = 2.0
    return para


def add_body_paragraph(doc, text):
    """Add a body paragraph with first-line indent."""
    para = doc.add_paragraph(text)
    para.style = doc.styles["Normal"]
    return para


def parse_markdown(md_text):
    """Parse markdown text into structured sections."""
    sections = []
    current_section = {"type": "body", "heading": None, "level": 0, "paragraphs": []}
    buffer = []

    for line in md_text.split("\n"):
        stripped = line.strip()

        # Detect headings
        heading_match = re.match(r"^(#{1,4})\s+(.+)$", stripped)
        if heading_match:
            # Flush buffer
            if buffer:
                text = " ".join(buffer).strip()
                if text:
                    current_section["paragraphs"].append(text)
                buffer = []
            # Save current section
            if current_section["heading"] or current_section["paragraphs"]:
                sections.append(current_section)
            level = len(heading_match.group(1))
            heading_text = heading_match.group(2).strip()
            current_section = {
                "type": "references" if "reference" in heading_text.lower() else "body",
                "heading": heading_text,
                "level": level,
                "paragraphs": [],
            }
            continue

        # Blank line = paragraph break
        if not stripped:
            if buffer:
                text = " ".join(buffer).strip()
                if text:
                    current_section["paragraphs"].append(text)
                buffer = []
            continue

        buffer.append(stripped)

    # Flush remaining
    if buffer:
        text = " ".join(buffer).strip()
        if text:
            current_section["paragraphs"].append(text)
    if current_section["heading"] or current_section["paragraphs"]:
        sections.append(current_section)

    return sections


def generate_docx(md_text, output_path, title=None):
    """Generate a Word document from markdown text."""
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    set_normal_style(doc)

    # Optional title
    if title:
        title_para = doc.add_paragraph(title)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.style = doc.styles["Normal"]
        title_para.paragraph_format.first_line_indent = Cm(0)
        for run in title_para.runs:
            run.bold = True
            run.font.size = Pt(16)
        doc.add_paragraph()  # Blank line after title

    sections = parse_markdown(md_text)

    for sec in sections:
        # Add heading
        if sec["heading"]:
            add_heading_styled(doc, sec["heading"], level=min(sec["level"], 3))

        # Add paragraphs
        if sec["type"] == "references":
            for para_text in sec["paragraphs"]:
                add_reference_paragraph(doc, para_text)
        else:
            for para_text in sec["paragraphs"]:
                add_body_paragraph(doc, para_text)

    doc.save(output_path)
    return output_path


def main():
    parser = argparse.ArgumentParser(description="Generate Word document from literature review markdown")
    parser.add_argument("--input", required=True, help="Input markdown file path")
    parser.add_argument("--output", required=True, help="Output .docx file path")
    parser.add_argument("--title", default=None, help="Optional document title")
    args = parser.parse_args()

    try:
        with open(args.input, "r", encoding="utf-8") as f:
            md_text = f.read()
    except FileNotFoundError:
        print(f"Error: Input file '{args.input}' not found.", file=sys.stderr)
        sys.exit(1)

    output = generate_docx(md_text, args.output, title=args.title)
    print(f"Word document generated: {output}")


if __name__ == "__main__":
    main()
