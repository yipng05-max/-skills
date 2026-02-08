#!/usr/bin/env python3
"""
Generate a formatted .docx file from manuscript content and style rules.

Usage:
    python format_docx.py <content.json> <style_rules.json> <output.docx>

The content.json should have the structure:
{
    "title": "Paper Title",
    "authors": ["Author One", "Author Two"],
    "affiliations": ["Affiliation 1", "Affiliation 2"],
    "abstract": "Abstract text...",
    "keywords": ["keyword1", "keyword2"],
    "sections": [
        {
            "heading": "Introduction",
            "level": 1,
            "content": ["Paragraph 1 text...", "Paragraph 2 text..."],
            "subsections": [
                {
                    "heading": "Background",
                    "level": 2,
                    "content": ["..."]
                }
            ]
        }
    ],
    "references": ["[1] Reference one...", "[2] Reference two..."],
    "footnotes": ["Footnote 1 text"],
    "acknowledgments": "Acknowledgment text..."
}

Dependencies:
    pip install python-docx
"""

import sys
import json
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm, Inches, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
except ImportError:
    print("Error: python-docx is required. Install with: pip install python-docx")
    sys.exit(1)


# Chinese font size mapping (字号 -> points)
CN_FONT_SIZE = {
    "初号": 42, "小初": 36, "一号": 26, "小一": 24,
    "二号": 22, "小二": 18, "三号": 16, "小三": 15,
    "四号": 14, "小四": 12, "五号": 10.5, "小五": 9,
    "六号": 7.5, "小六": 6.5, "七号": 5.5, "八号": 5,
}


def resolve_font_size(size_spec):
    """Resolve a font size specification to points."""
    if size_spec is None:
        return None
    if isinstance(size_spec, (int, float)):
        return size_spec
    if isinstance(size_spec, str):
        # Try Chinese font size name
        if size_spec in CN_FONT_SIZE:
            return CN_FONT_SIZE[size_spec]
        # Try parsing "12pt" or "12"
        size_spec = size_spec.replace("pt", "").strip()
        try:
            return float(size_spec)
        except ValueError:
            return None
    return None


def resolve_alignment(align_str):
    """Convert alignment string to enum."""
    mapping = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        "distribute": WD_ALIGN_PARAGRAPH.DISTRIBUTE,
    }
    return mapping.get(align_str, WD_ALIGN_PARAGRAPH.JUSTIFY)


def set_paragraph_format(paragraph, fmt):
    """Apply formatting to a paragraph from a format dict."""
    pf = paragraph.paragraph_format

    if fmt.get("alignment"):
        pf.alignment = resolve_alignment(fmt["alignment"])

    if fmt.get("first_line_indent_cm") is not None:
        pf.first_line_indent = Cm(fmt["first_line_indent_cm"])

    if fmt.get("left_indent_cm") is not None:
        pf.left_indent = Cm(fmt["left_indent_cm"])

    if fmt.get("space_before_pt") is not None:
        pf.space_before = Pt(fmt["space_before_pt"])

    if fmt.get("space_after_pt") is not None:
        pf.space_after = Pt(fmt["space_after_pt"])

    # Line spacing
    line_spacing = fmt.get("line_spacing")
    rule = fmt.get("line_spacing_rule")
    if line_spacing is not None:
        if isinstance(line_spacing, str) and line_spacing.endswith("pt"):
            # Exact point spacing
            pf.line_spacing = Pt(float(line_spacing.replace("pt", "")))
            pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        elif rule == "EXACTLY":
            pf.line_spacing = Pt(float(str(line_spacing).replace("pt", "")))
            pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        elif rule == "AT_LEAST":
            pf.line_spacing = Pt(float(str(line_spacing).replace("pt", "")))
            pf.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
        else:
            # Multiple (1.0, 1.5, 2.0, etc.)
            try:
                pf.line_spacing = float(line_spacing)
                pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            except (ValueError, TypeError):
                pass


def set_run_format(run, fmt):
    """Apply character formatting to a run."""
    font = run.font

    font_names = fmt.get("font_names", [])
    font_name = font_names[0] if font_names else fmt.get("font_name")
    if font_name:
        font.name = font_name
        # Set East Asian font for Chinese support
        r = run._element
        rPr = r.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
            rPr.insert(0, rFonts)
        # If font is Chinese, set as East Asian font
        if any("\u4e00" <= c <= "\u9fff" for c in font_name):
            rFonts.set(qn("w:eastAsia"), font_name)
        else:
            rFonts.set(qn("w:ascii"), font_name)
            rFonts.set(qn("w:hAnsi"), font_name)

    font_sizes = fmt.get("font_sizes_pt", [])
    size = font_sizes[0] if font_sizes else fmt.get("font_size_pt")
    size = resolve_font_size(size)
    if size:
        font.size = Pt(size)

    if fmt.get("bold") is not None:
        font.bold = fmt["bold"]

    if fmt.get("italic") is not None:
        font.italic = fmt["italic"]

    if fmt.get("color"):
        try:
            font.color.rgb = RGBColor.from_string(fmt["color"])
        except (ValueError, AttributeError):
            pass


def add_formatted_paragraph(doc, text, role_fmt, style_name=None):
    """Add a paragraph with formatting based on role format dict."""
    if style_name:
        try:
            para = doc.add_paragraph(style=style_name)
        except KeyError:
            para = doc.add_paragraph()
    else:
        para = doc.add_paragraph()

    run = para.add_run(text)
    set_paragraph_format(para, role_fmt)
    set_run_format(run, role_fmt)
    return para


def setup_page(doc, layout):
    """Configure page layout from style rules."""
    if not layout:
        return

    # Use first section's layout
    sec_info = layout[0] if isinstance(layout, list) else layout
    section = doc.sections[0]

    # Page size
    if sec_info.get("page_width_cm"):
        section.page_width = Cm(sec_info["page_width_cm"])
    if sec_info.get("page_height_cm"):
        section.page_height = Cm(sec_info["page_height_cm"])

    # Orientation
    if sec_info.get("orientation") == "landscape":
        section.orientation = WD_ORIENT.LANDSCAPE
        # Swap width and height if needed
        if section.page_width < section.page_height:
            section.page_width, section.page_height = section.page_height, section.page_width

    # Margins
    margins = sec_info.get("margins", {})
    if margins.get("top_cm") is not None:
        section.top_margin = Cm(margins["top_cm"])
    if margins.get("bottom_cm") is not None:
        section.bottom_margin = Cm(margins["bottom_cm"])
    if margins.get("left_cm") is not None:
        section.left_margin = Cm(margins["left_cm"])
    if margins.get("right_cm") is not None:
        section.right_margin = Cm(margins["right_cm"])

    # Header/footer distance
    if sec_info.get("header_distance_cm"):
        section.header_distance = Cm(sec_info["header_distance_cm"])
    if sec_info.get("footer_distance_cm"):
        section.footer_distance = Cm(sec_info["footer_distance_cm"])

    # Header text
    if sec_info.get("header_text"):
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.text = sec_info["header_text"]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Footer text
    if sec_info.get("footer_text"):
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.text = sec_info["footer_text"]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER


def get_role_format(style_rules, role, defaults=None):
    """Get formatting for a specific role from style rules."""
    fmt = {}
    if defaults:
        fmt.update(defaults)

    para_fmt = style_rules.get("paragraph_formatting", {})
    if role in para_fmt:
        role_data = para_fmt[role]
        for key in ["font_names", "font_sizes_pt", "bold", "italic", "alignment",
                     "line_spacing", "line_spacing_rule", "first_line_indent_cm",
                     "space_before_pt", "space_after_pt"]:
            if role_data.get(key) is not None:
                fmt[key] = role_data[key]

    return fmt


def build_default_formats():
    """Build sensible default formatting for each role."""
    return {
        "title": {
            "font_names": ["Times New Roman"],
            "font_sizes_pt": [16],
            "bold": True,
            "alignment": "center",
            "space_after_pt": 12,
        },
        "authors": {
            "font_names": ["Times New Roman"],
            "font_sizes_pt": [12],
            "alignment": "center",
            "space_after_pt": 6,
        },
        "affiliations": {
            "font_names": ["Times New Roman"],
            "font_sizes_pt": [10],
            "italic": True,
            "alignment": "center",
            "space_after_pt": 12,
        },
        "abstract_label": {
            "font_names": ["Times New Roman"],
            "font_sizes_pt": [12],
            "bold": True,
            "alignment": "left",
        },
        "abstract": {
            "font_names": ["Times New Roman"],
            "font_sizes_pt": [10],
            "alignment": "justify",
            "first_line_indent_cm": 0,
            "space_after_pt": 12,
        },
        "keywords": {
            "font_names": ["Times New Roman"],
            "font_sizes_pt": [10],
            "alignment": "left",
            "space_after_pt": 12,
        },
        "heading_1": {
            "font_names": ["Times New Roman"],
            "font_sizes_pt": [14],
            "bold": True,
            "alignment": "left",
            "space_before_pt": 18,
            "space_after_pt": 12,
        },
        "heading_2": {
            "font_names": ["Times New Roman"],
            "font_sizes_pt": [12],
            "bold": True,
            "alignment": "left",
            "space_before_pt": 12,
            "space_after_pt": 6,
        },
        "heading_3": {
            "font_names": ["Times New Roman"],
            "font_sizes_pt": [12],
            "bold": False,
            "italic": True,
            "alignment": "left",
            "space_before_pt": 6,
            "space_after_pt": 6,
        },
        "body": {
            "font_names": ["Times New Roman"],
            "font_sizes_pt": [12],
            "alignment": "justify",
            "line_spacing": 1.5,
            "first_line_indent_cm": 0.75,
            "space_after_pt": 0,
        },
        "caption": {
            "font_names": ["Times New Roman"],
            "font_sizes_pt": [10],
            "alignment": "center",
            "bold": False,
            "space_before_pt": 6,
            "space_after_pt": 6,
        },
        "reference": {
            "font_names": ["Times New Roman"],
            "font_sizes_pt": [10],
            "alignment": "justify",
            "first_line_indent_cm": 0,
            "left_indent_cm": 0,
            "space_after_pt": 3,
        },
    }


def add_section_content(doc, section_data, style_rules, defaults, numbering_prefix=""):
    """Recursively add section headings and content."""
    level = section_data.get("level", 1)
    heading_text = section_data.get("heading", "")

    # Add numbering prefix if present
    if numbering_prefix and heading_text:
        heading_text = f"{numbering_prefix} {heading_text}"

    # Add heading
    role = f"heading_{level}"
    fmt = get_role_format(style_rules, role, defaults.get(role))
    add_formatted_paragraph(doc, heading_text, fmt)

    # Add content paragraphs
    body_fmt = get_role_format(style_rules, "body", defaults.get("body"))
    for para_text in section_data.get("content", []):
        if para_text.strip():
            add_formatted_paragraph(doc, para_text, body_fmt)

    # Add subsections
    for i, subsec in enumerate(section_data.get("subsections", []), 1):
        sub_prefix = f"{numbering_prefix}.{i}" if numbering_prefix else ""
        add_section_content(doc, subsec, style_rules, defaults, sub_prefix)


def format_document(content, style_rules, output_path):
    """Generate a formatted .docx from content and style rules."""
    doc = Document()
    defaults = build_default_formats()

    # Page setup
    setup_page(doc, style_rules.get("page_layout"))

    # Title
    title_fmt = get_role_format(style_rules, "title", defaults["title"])
    if content.get("title"):
        add_formatted_paragraph(doc, content["title"], title_fmt)

    # Authors
    if content.get("authors"):
        auth_fmt = get_role_format(style_rules, "authors", defaults["authors"])
        authors_text = ", ".join(content["authors"])
        add_formatted_paragraph(doc, authors_text, auth_fmt)

    # Affiliations
    if content.get("affiliations"):
        aff_fmt = get_role_format(style_rules, "affiliations", defaults["affiliations"])
        for aff in content["affiliations"]:
            add_formatted_paragraph(doc, aff, aff_fmt)

    # Abstract
    if content.get("abstract"):
        label_fmt = get_role_format(style_rules, "abstract_label", defaults["abstract_label"])
        add_formatted_paragraph(doc, "Abstract", label_fmt)
        abs_fmt = get_role_format(style_rules, "abstract", defaults["abstract"])
        add_formatted_paragraph(doc, content["abstract"], abs_fmt)

    # Keywords
    if content.get("keywords"):
        kw_fmt = get_role_format(style_rules, "keywords", defaults["keywords"])
        kw_text = "Keywords: " + "; ".join(content["keywords"])
        add_formatted_paragraph(doc, kw_text, kw_fmt)

    # Sections
    for i, section in enumerate(content.get("sections", []), 1):
        add_section_content(doc, section, style_rules, defaults, str(i))

    # Acknowledgments
    if content.get("acknowledgments"):
        ack_fmt = get_role_format(style_rules, "heading_1", defaults["heading_1"])
        add_formatted_paragraph(doc, "Acknowledgments", ack_fmt)
        body_fmt = get_role_format(style_rules, "body", defaults["body"])
        add_formatted_paragraph(doc, content["acknowledgments"], body_fmt)

    # References
    if content.get("references"):
        ref_heading_fmt = get_role_format(style_rules, "heading_1", defaults["heading_1"])
        add_formatted_paragraph(doc, "References", ref_heading_fmt)
        ref_fmt = get_role_format(style_rules, "reference", defaults["reference"])
        for ref in content["references"]:
            add_formatted_paragraph(doc, ref, ref_fmt)

    doc.save(str(output_path))
    print(f"Formatted document saved to: {output_path}")


def main():
    if len(sys.argv) < 4:
        print(f"Usage: {sys.argv[0]} <content.json> <style_rules.json> <output.docx>")
        sys.exit(1)

    content_path = Path(sys.argv[1])
    rules_path = Path(sys.argv[2])
    output_path = Path(sys.argv[3])

    with open(content_path, "r", encoding="utf-8") as f:
        content = json.load(f)

    with open(rules_path, "r", encoding="utf-8") as f:
        style_rules = json.load(f)

    format_document(content, style_rules, output_path)


if __name__ == "__main__":
    main()
